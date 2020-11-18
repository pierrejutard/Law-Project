from docx import Document
from docx.shared import Inches

# function creating a Word type contrat file
def Create_Word_Contract_File(Result_dico):
    doc = Document()
    for i_part_text in Result_dico.values():
        p = doc.add_paragraph('')
        i = 0
        while i != len(i_part_text):
            carac = i_part_text[i]
            #underline case
            if carac =='¤':
                j = i+1
                Next_carac = i_part_text[j]
                #bold case
                if Next_carac =='*':
                    j = j+1
                    Next_carac = i_part_text[j]
                    while Next_carac !='*':
                        runner = p.add_run(Next_carac)
                        runner.bold = True
                        runner.underline = True
                        j = j+1
                        Next_carac = i_part_text[j]
                    i = j+2
                else:
                    #underline case
                    while Next_carac !='¤':
                        runner = p.add_run(Next_carac)
                        runner.underline = True
                        j = j+1
                        Next_carac = i_part_text[j]
                    i = j+1
            elif carac =='*':
                #bold case
                j = i + 1
                Next_carac = i_part_text[j]
                while Next_carac !='*':
                    runner = p.add_run(Next_carac)
                    runner.bold = True
                    j = j+1
                    Next_carac = i_part_text[j]
                i = j+1
            else:
                p.add_run(carac)
                i = i+1
    return doc